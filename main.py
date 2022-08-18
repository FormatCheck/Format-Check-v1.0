# import modules
from docx import Document  # for accessing the document by python-docx
import docx2txt as docx  # for accessing the document by docx2txt
import streamlit as st  # for web app
from PIL import Image  # for logo
import re  # for regex pattern used in reference counter

# configure web app logo and name
logo = Image.open('logo.png')
st.set_page_config(page_title='Format Check v1.0', page_icon=logo, layout='wide')


# font name program function
def font_name():
    # TODO start of Font Name code --------------------
    # add font name program banner
    st.subheader("**Font Name**")
    # access Word document file
    WordFile = Document(uploaded_file)

    # TODO Font Name code for Title style text
    # check font name for Title style text (text on title page) ********************
    title_font = set()  # store all Title style font names in the set title_font
    title_wrong_font = set()  # store unacceptable Title style font names in the set title_wrong_font
    title_wrong_font_words = sorted(set())  # store Title text that are in unacceptable fonts in the sorted list title_wrong_font_words
    CORRECT_FONT_NAME_TITLE = 'Times New Roman'  # state the specified font for Title style and store in the variable CORRECT_FONT_NAME_TITLE
    for paragraph in WordFile.paragraphs:
        if 'Title' == paragraph.style.name:
            for run in paragraph.runs:
                # add fonts from each run into the set title_font
                title_font.add(run.font.name)
                # check if fonts are unacceptable, if so, store in the set title_wrong_font
                if run.font.name != CORRECT_FONT_NAME_TITLE:
                    title_wrong_font.add(run.font.name)
                    # append Title text that contain unacceptable fonts in the sorted list title_wrong_font_words
                    title_wrong_font_words.append(run.text)

        # check if all elements in title_font are not CORRECT_FONT_NAME_TITLE and if title_font is not empty
    if {CORRECT_FONT_NAME_TITLE} != title_font and len(title_font) != 0:
        # print this if all elements in title_font are not CORRECT_FONT_NAME_TITLE and if title_font is not empty and print title_wrong_font and title_wrong_font_words contents
        st.error(f'''
        ❌ Title style text have incorrect font(s): {', '.join(map(str, title_wrong_font))}  
        🡆 Incorrect font(s) in Title style text found here: {' >> '.join(map(str, title_wrong_font_words))}
        ''')
        # check if title_font is empty, if so this means that Title style was not found
    elif len(title_font) == 0:
        # print this if title_font is empty, since Title style was not found
        st.info("ℹ️Title style text font name not found as this style was not used.")

    # TODO Font Name code for Heading 1 style text
    # check font name for Heading 1 style text (main headings) ********************
    h1_font = set()  # store all Heading 1 style font names in the set h1_font
    h1_wrong_font = set()  # store unacceptable Heading 1 style font names in the set h1_wrong_font
    h1_wrong_font_words = sorted(set())  # store Heading 1 text that are in unacceptable fonts in the sorted list h1_wrong_font_words
    CORRECT_FONT_NAME_H1 = 'Times New Roman'  # state the specified font for Heading 1 style and store in the variable CORRECT_FONT_NAME_H1
    for paragraph in WordFile.paragraphs:
        if 'Heading 1' == paragraph.style.name:
            for run in paragraph.runs:
                # add fonts from each run into the set h1_font
                h1_font.add(run.font.name)
                # check if fonts are unacceptable, if so, store in the set h1_wrong_font
                if run.font.name != CORRECT_FONT_NAME_H1:
                    h1_wrong_font.add(run.font.name)
                    # append Heading 1 text that contain unacceptable fonts in the sorted list h1_wrong_font_words
                    h1_wrong_font_words.append(run.text)

        # check if all elements in h1_font are not CORRECT_FONT_NAME_H1 and if h1_font is not empty
    if {CORRECT_FONT_NAME_H1} != h1_font and len(h1_font) != 0:
        # print this if all elements in h1_font are not CORRECT_FONT_NAME_H1 and if h1_font is not empty and print h1_wrong_font and h1_wrong_font_words contents
        st.error(f'''
        ❌ Heading 1 style text have incorrect font(s): {', '.join(map(str, h1_wrong_font))}  
        🡆 Incorrect font(s) in Heading 1 style text found here: {' >> '.join(map(str, h1_wrong_font_words))}
        ''')
        # check if h1_font is empty, if so this means that Heading 1 style was not found
    elif len(h1_font) == 0:
        # print this if h1_font is empty, since Heading 1 style was not found
        st.info("ℹ️Heading 1 style text font name not found as this style was not used.")

    # TODO Font Name code for Heading 2 style text
    # check font name for Heading 2 style text (sub headings) ********************
    h2_font = set()  # store all Heading 2 style font names in the set h2_font
    h2_wrong_font = set()  # store unacceptable Heading 2 style font names in the set h2_wrong_font
    h2_wrong_font_words = sorted(set())  # store Heading 2 text that are in unacceptable fonts in the sorted list h2_wrong_font_words
    CORRECT_FONT_NAME_H2 = 'Times New Roman'  # state the specified font for Heading 2 style and store in the variable CORRECT_FONT_NAME_H2
    for paragraph in WordFile.paragraphs:
        if 'Heading 2' == paragraph.style.name:
            for run in paragraph.runs:
                # add fonts from each run into the set h2_font
                h2_font.add(run.font.name)
                # check if fonts are unacceptable, if so, store in the set h2_wrong_font
                if run.font.name != CORRECT_FONT_NAME_H2:
                    h2_wrong_font.add(run.font.name)
                    # append Heading 2 text that contain unacceptable fonts in the sorted list h2_wrong_font_words
                    h2_wrong_font_words.append(run.text)

        # check if all elements in h2_font are not CORRECT_FONT_NAME_H2 and if h2_font is not empty
    if {CORRECT_FONT_NAME_H2} != h2_font and len(h2_font) != 0:
        # print this if all elements in h2_font are not CORRECT_FONT_NAME_H2 and if h2_font is not empty and print h2_wrong_font and h2_wrong_font_words contents
        st.error(f'''
        ❌ Heading 2 style text have incorrect font(s): {', '.join(map(str, h2_wrong_font))}  
        🡆 Incorrect font(s) in Heading 2 style text found here: {' >> '.join(map(str, h2_wrong_font_words))}
        ''')
        # check if h2_font is empty, if so this means that Heading 2 style was not found
    elif len(h2_font) == 0:
        # print this if h2_font is empty, since Heading 2 style was not found
        st.info("ℹ️Heading 2 style text font name not found as this style was not used.")

    # TODO Font Name code for Heading 3 style text
    # check font name for Heading 3 style text (sub headings) ********************
    h3_font = set()  # store all Heading 3 style font names in the set h3_font
    h3_wrong_font = set()  # store unacceptable Heading 3 style font names in the set h3_wrong_font
    h3_wrong_font_words = sorted(set())  # store Heading 3 text that are in unacceptable fonts in the sorted list h3_wrong_font_words
    CORRECT_FONT_NAME_H3 = 'Times New Roman'  # state the specified font for Heading 3 style and store in the variable CORRECT_FONT_NAME_H3
    for paragraph in WordFile.paragraphs:
        if 'Heading 3' == paragraph.style.name:
            for run in paragraph.runs:
                # add fonts from each run into the set h3_font
                h3_font.add(run.font.name)
                # check if fonts are unacceptable, if so, store in the set h3_wrong_font
                if run.font.name != CORRECT_FONT_NAME_H3:
                    h3_wrong_font.add(run.font.name)
                    # append Heading 3 text that contain unacceptable fonts in the sorted list h3_wrong_font_words
                    h3_wrong_font_words.append(run.text)

        # check if all elements in h3_font are not CORRECT_FONT_NAME_H3 and if h3_font is not empty
    if {CORRECT_FONT_NAME_H3} != h3_font and len(h3_font) != 0:
        # print this if all elements in h3_font are not CORRECT_FONT_NAME_H3 and if h3_font is not empty and print h3_wrong_font and h3_wrong_font_words contents
        st.error(f'''
        ❌ Heading 3 style text have incorrect font(s): {', '.join(map(str, h3_wrong_font))}  
        🡆 Incorrect font(s) in Heading 3 style text found here: {' >> '.join(map(str, h3_wrong_font_words))}
        ''')
        # check if h3_font is empty, if so this means that Heading 3 style was not found
    elif len(h3_font) == 0:
        # print this if h3_font is empty, since Heading 3 style was not found
        st.info("ℹ️Heading 3 style text font name not found as this style was not used.")

    # TODO Font Name code for Heading 4 style text
    # check font name for Heading 4 style text (sub headings) ********************
    h4_font = set()  # store all Heading 4 style font names in the set h4_font
    h4_wrong_font = set()  # store unacceptable Heading 4 style font names in the set h4_wrong_font
    h4_wrong_font_words = sorted(set())  # store Heading 4 text that are in unacceptable fonts in the sorted list h4_wrong_font_words
    CORRECT_FONT_NAME_H4 = 'Times New Roman'  # state the specified font for Heading 4 style and store in the variable CORRECT_FONT_NAME_H4
    for paragraph in WordFile.paragraphs:
        if 'Heading 4' == paragraph.style.name:
            for run in paragraph.runs:
                # add fonts from each run into the set h4_font
                h4_font.add(run.font.name)
                # check if fonts are unacceptable, if so, store in the set h4_wrong_font
                if run.font.name != CORRECT_FONT_NAME_H4:
                    h4_wrong_font.add(run.font.name)
                    # append Heading 4 text that contain unacceptable fonts in the sorted list h4_wrong_font_words
                    h4_wrong_font_words.append(run.text)

        # check if all elements in h4_font are not CORRECT_FONT_NAME_H4 and if h4_font is not empty
    if {CORRECT_FONT_NAME_H4} != h4_font and len(h4_font) != 0:
        # print this if all elements in h4_font are not CORRECT_FONT_NAME_H4 and if h4_font is not empty and print h4_wrong_font and h4_wrong_font_words contents
        st.error(f'''
        ❌ Heading 4 style text have incorrect font(s): {', '.join(map(str, h4_wrong_font))}  
        🡆 Incorrect font(s) in Heading 4 style text found here: {' >> '.join(map(str, h4_wrong_font_words))}
        ''')
        # check if h4_font is empty, if so this means that Heading 4 style was not found
    elif len(h4_font) == 0:
        # print this if h4_font is empty, since Heading 4 style was not found
        st.info("ℹ️Heading 4 style text font name not found as this style was not used.")

    # TODO Font Name code for Normal style text
    # check font name for Normal style text (body text/paragraphs) ********************
    norm_font = set()  # store all Normal style font names in the set norm_font
    norm_wrong_font = set()  # store unacceptable Normal style font names in the set norm_wrong_font
    norm_wrong_font_words = sorted(set())  # store Normal text that are in unacceptable fonts in the sorted list norm_wrong_font_words
    CORRECT_FONT_NAME_NORM = 'Times New Roman'  # state the specified font for Normal style and store in the variable CORRECT_FONT_NAME_NORM
    for paragraphs in WordFile.paragraphs:
        if 'Normal' == paragraphs.style.name:
            for run in paragraphs.runs:
                # add fonts from each run into the set norm_font
                norm_font.add(run.font.name)
                # check if fonts are unacceptable, if so, store in the set norm_wrong_font
                if run.font.name != CORRECT_FONT_NAME_NORM:
                    norm_wrong_font.add(run.font.name)
                    # append Normal text that contain unacceptable fonts in the sorted list norm_wrong_font_words
                    norm_wrong_font_words.append(run.text)

        # check if all elements in norm_font are not CORRECT_FONT_NAME_NORM
    if {CORRECT_FONT_NAME_NORM} != norm_font and len(norm_font) != 0:
        # print this if all elements in norm_font are not CORRECT_FONT_NAME_NORM and print norm_wrong_font and norm_wrong_font_words contents
        st.error(f'''
        ❌ Normal style text have incorrect font(s): {', '.join(map(str, norm_wrong_font))}  
        🡆 Incorrect font(s) in Normal style text found here: {' >> '.join(map(str, norm_wrong_font_words))}
        ''')
    elif len(norm_font) == 0:
        # print this if norm_font is empty, which means that Normal style was not found
        st.info("ℹ️Normal style text font name not found as this style was not used.")

    # TODO Font Name code for List Paragraph style text
    # check font name for List Paragraph style text (bullet list) ********************
    list_font = set()  # store all List Paragraph style font names in the set list_font
    list_wrong_font = set()  # store unacceptable List Paragraph style font names in the set list_wrong_font
    list_wrong_font_words = sorted(set())  # store List Paragraph text that are in unacceptable fonts in the sorted list cap_wrong_font_words
    CORRECT_FONT_NAME_LIST = 'Times New Roman'  # state the specified font for List Paragraph style and store in the variable CORRECT_FONT_NAME_LIST
    for paragraph in WordFile.paragraphs:
        if 'List Paragraph' == paragraph.style.name:
            for run in paragraph.runs:
                # add fonts from each run into the set list_font
                list_font.add(run.font.name)
                # check if fonts are unacceptable, if so, store in the set list_wrong_font
                if run.font.name != CORRECT_FONT_NAME_LIST:
                    list_wrong_font.add(run.font.name)
                    # append List Paragraph text that contain unacceptable fonts in the sorted list list_wrong_font_words
                    list_wrong_font_words.append(run.text)

        # check if all elements in list_font are not CORRECT_FONT_NAME_LIST and if list_font is not empty
    if {CORRECT_FONT_NAME_LIST} != list_font and len(list_font) != 0:
        # print this if all elements in list_font are not CORRECT_FONT_NAME_LIST and if list_font is not empty and print list_wrong_font and list_wrong_font_words contents
        st.error(f'''
        ❌ List Paragraph style text have incorrect font(s): {', '.join(map(str, list_wrong_font))}  
        🡆 Incorrect font(s) in List Paragraph style text found here: {' >> '.join(map(str, list_wrong_font_words))}
        ''')
        # check if list_font is empty, if so this means that List Paragraph style was not found
    elif len(list_font) == 0:
        # print this if list_font is empty, since List Paragraph style was not found
        st.info("ℹ️List Paragraph style text font name not found as this style was not used.")

    # TODO Font Name code for Caption style text
    # check font name for Caption style text (figure and table captions) ********************
    cap_font = set()  # store all Caption style font names in the set cap_font
    cap_wrong_font = set()  # store unacceptable Caption style font names in the set cap_wrong_font
    cap_wrong_font_words = sorted(set())  # store Caption text that are in unacceptable fonts in the sorted list cap_wrong_font_words
    CORRECT_FONT_NAME_CAP = 'Times New Roman'  # state the specified font for Caption style and store in the variable CORRECT_FONT_NAME_CAP
    for paragraph in WordFile.paragraphs:
        if 'Caption' == paragraph.style.name:
            for run in paragraph.runs:
                # add fonts from each run into the set cap_font
                cap_font.add(run.font.name)
                # check if fonts are unacceptable, if so, store in the set cap_wrong_font
                if run.font.name != CORRECT_FONT_NAME_CAP:
                    cap_wrong_font.add(run.font.name)
                    # append Caption text that contain unacceptable fonts in the sorted list norm_wrong_font_words
                    cap_wrong_font_words.append(run.text)

        # check if all elements in cap_font are not CORRECT_FONT_NAME_CAP and if cap_font is not empty
    if {CORRECT_FONT_NAME_CAP} != cap_font and len(cap_font) != 0:
        # print this if all elements in cap_font are not CORRECT_FONT_NAME_CAP and if cap_font is not empty and print cap_wrong_font and cap_wrong_font_words contents
        st.error(f'''
        ❌ Caption style text have incorrect font(s): {', '.join(map(str, cap_wrong_font))}  
        🡆 Incorrect font(s) in Caption style text found here: {' >> '.join(map(str, cap_wrong_font_words))}
        ''')
        # check if cap_font is empty, if so this means that Caption style was not found
    elif len(cap_font) == 0:
        # print this if cap_font is empty, since Caption style was not found
        st.info("ℹ️Caption style text font name not found as this style was not used.")


# font size program function
def font_size():
    # TODO start of Font Size code --------------------
    # add font size program banner
    st.subheader("**Font Size**")
    # access Word document file
    WordFile = Document(uploaded_file)

    # TODO Font Size code for Title style text
    # check font size for Title style text (text on title page) ********************
    title_size = set()  # store all Title style font sizes in the set title_size
    title_wrong_size = set()  # store unacceptable Title style font sizes in the set title_wrong_size
    title_wrong_size_words = sorted(set())  # store Title text that are in unacceptable font size in the sorted list title_wrong_size_words
    CORRECT_FONT_SIZE_TITLE = 381000  # state the specified font size for Title style and store in the variable CORRECT_FONT_SIZE_TITLE; 38100/12700 = 30pt
    for paragraph in WordFile.paragraphs:
        if 'Title' == paragraph.style.name:
            for run in paragraph.runs:
                # add fonts from each run into the set title_size
                title_size.add(run.font.size)
                # check if font sizes are unacceptable, if so, store in the set title_wrong_size
                if run.font.size != CORRECT_FONT_SIZE_TITLE:
                    if run.font.size is not None:
                        title_wrong_size.add(run.font.size / 12700)
                    else:
                        title_wrong_size.add(run.font.size)
                    # append Title text that contain unacceptable font sizes in the sorted list title_wrong_size_words
                    title_wrong_size_words.append(run.text)

        # check if all elements in title_size are not CORRECT_FONT_SIZE_TITLE and if title_size is not empty
    if {CORRECT_FONT_SIZE_TITLE} != title_size and len(title_size) != 0:
        # print this if all elements in title_size are not CORRECT_FONT_SIZE_TITLE and if title_size is not empty and print title_wrong_size and title_wrong_size_words contents
        st.error(f'''
        ❌ Title style text have incorrect font sizes(s): {', '.join(map(str, title_wrong_size))}  
        🡆 Incorrect font size(s) in Title style text found here: {' >> '.join(map(str, title_wrong_size_words))}
        ''')
        # check if title_size is empty, if so this means that Title style was not found
    elif len(title_size) == 0:
        # print this if title_size is empty, since Title style was not found
        st.info("ℹ️Title style text font size not found as this style was not used.")

    # TODO Font Size code for Heading 1 style text
    # check font size for Heading 1 style text (main headings) ********************
    h1_size = set()  # store all Heading 1 style font sizes in the set h1_size
    h1_wrong_size = set()  # store unacceptable Heading 1 style font sizes in the set h1_wrong_size
    h1_wrong_size_words = sorted(set())  # store Heading 1 text that are in unacceptable font size in the sorted list h1_wrong_size_words
    CORRECT_FONT_SIZE_H1 = 177800  # state the specified font size for Heading 1 style and store in the variable CORRECT_FONT_SIZE_H1; 177800/12700 = 14pt
    for paragraph in WordFile.paragraphs:
        if 'Heading 1' == paragraph.style.name:
            for run in paragraph.runs:
                # add font sizes from each run into the set norm_size
                h1_size.add(run.font.size)
                # check if font sizes are unacceptable, if so, store in the set h1_wrong_size
                if run.font.size != CORRECT_FONT_SIZE_H1:
                    if run.font.size is not None:
                        h1_wrong_size.add(run.font.size / 12700)
                    else:
                        h1_wrong_size.add(run.font.size)
                    # append Heading 1 text that contain unacceptable font sizes in the sorted list h1_wrong_size_words
                    h1_wrong_size_words.append(run.text)

        # check if all elements in h1_size are not CORRECT_FONT_SIZE_H1 and if h1_size is not empty
    if {CORRECT_FONT_SIZE_H1} != h1_size and len(h1_size) != 0:
        # print this if all elements in h1_size are not CORRECT_FONT_SIZE_H1 and if h1_size is not empty and print h1_wrong_size and h1_wrong_size_words contents
        st.error(f'''
        ❌ Heading 1 style text have incorrect font sizes(s): {', '.join(map(str, h1_wrong_size))}  
        🡆 Incorrect font size(s) in Heading 1 style text found here: {' >> '.join(map(str, h1_wrong_size_words))}
        ''')
        # check if h1_size is empty, if so this means that Heading 1 style was not found
    elif len(h1_size) == 0:
        # print this if h1_size is empty, since Heading 1 style was not found
        st.info("ℹ️Heading 1 style text font size not found as this style was not used.")

    # TODO Font Size code for Heading 2 style text
    # check font size for Heading 2 style text (sub headings) ********************
    h2_size = set()  # store all Heading 2 style font sizes in the set h2_size
    h2_wrong_size = set()  # store unacceptable Heading 2 style font sizes in the set h2_wrong_size
    h2_wrong_size_words = sorted(set())  # store Heading 2 text that are in unacceptable font size in the sorted list h2_wrong_size_words
    CORRECT_FONT_SIZE_H2 = 177800  # state the specified font size for Heading 2 style and store in the variable CORRECT_FONT_SIZE_H2; 177800/12700 = 14pt
    for paragraph in WordFile.paragraphs:
        if 'Heading 2' == paragraph.style.name:
            for run in paragraph.runs:
                # add font sizes from each run into the set h2_size
                h2_size.add(run.font.size)
                # check if font sizes are unacceptable, if so, store in the set h2_wrong_size
                if run.font.size != CORRECT_FONT_SIZE_H2:
                    if run.font.size is not None:
                        h2_wrong_size.add(run.font.size / 12700)
                    else:
                        h2_wrong_size.add(run.font.size)
                    # append Heading 2 text that contain unacceptable font sizes in the sorted list h2_wrong_size_words
                    h2_wrong_size_words.append(run.text)

        # check if all elements in h2_size are not CORRECT_FONT_SIZE_H2 and if h2_size is not empty
    if {CORRECT_FONT_SIZE_H2} != h2_size and len(h2_size) != 0:
        # print this if all elements in h2_size are not CORRECT_FONT_SIZE_H2 and if h2_size is not empty and print h2_wrong_size and h2_wrong_size_words contents
        st.error(f'''
        ❌ Heading 2 style text have incorrect font sizes(s): {', '.join(map(str, h2_wrong_size))}  
        🡆 Incorrect font size(s) in Heading 2 style text found here: {' >> '.join(map(str, h2_wrong_size_words))}
        ''')
        # check if h2_size is empty, if so this means that Heading 2 style was not found
    elif len(h2_size) == 0:
        # print this if h2_size is empty, since Heading 2 style was not found
        st.info("ℹ️Heading 2 style text font size not found as this style was not used.")

    # TODO Font Size code for Heading 3 style text
    # check font size for Heading 3 style text (sub headings) ********************
    h3_size = set()  # store all Heading 3 style font sizes in the set h3_size
    h3_wrong_size = set()  # store unacceptable Heading 3 style font sizes in the set h3_wrong_size
    h3_wrong_size_words = sorted(set())  # store Heading 3 text that are in unacceptable font size in the sorted list h3_wrong_size_words
    CORRECT_FONT_SIZE_H3 = 177800  # # state the specified font size for Heading 3 style and store in the variable CORRECT_FONT_SIZE_H3; 177800/12700 = 14pt
    for paragraph in WordFile.paragraphs:
        if 'Heading 3' == paragraph.style.name:
            for run in paragraph.runs:
                # add font sizes from each run into the set h3_size
                h3_size.add(run.font.size)
                # check if font sizes are unacceptable, if so, store in the set h3_wrong_size
                if run.font.size != CORRECT_FONT_SIZE_H3:
                    if run.font.size is not None:
                        h3_wrong_size.add(run.font.size / 12700)
                    else:
                        h3_wrong_size.add(run.font.size)
                    # append Heading 3 text that contain unacceptable font sizes in the sorted list h3_wrong_size_words
                    h3_wrong_size_words.append(run.text)

        # check if all elements in h3_size are not CORRECT_FONT_SIZE_H3 and if h3_size is not empty
    if {CORRECT_FONT_SIZE_H3} != h3_size and len(h3_size) != 0:
        # print this if all elements in h3_size are not CORRECT_FONT_SIZE_H3 and if h3_size is not empty and print h3_wrong_size and h3_wrong_size_words contents
        st.error(f'''
        ❌ Heading 3 style text have incorrect font sizes(s): {', '.join(map(str, h3_wrong_size))}  
        🡆 Incorrect font size(s) in Heading 3 style text found here: {' >> '.join(map(str, h3_wrong_size_words))}
        ''')
        # check if h3_size is empty, if so this means that Heading 3 style was not found
    elif len(h3_size) == 0:
        # print this if h3_size is empty, since Heading 3 style was not found
        st.info("ℹ️Heading 3 style text font size not found as this style was not used.")

    # TODO Font Size code for Heading 4 style text
    # check font size for Heading 4 style text (sub headings) ********************
    h4_size = set()  # store all Heading 4 style font sizes in the set h4_size
    h4_wrong_size = set()  # store unacceptable Heading 4 style font sizes in the set h4_wrong_size
    h4_wrong_size_words = sorted(set())  # store Heading 4 text that are in unacceptable font size in the sorted list h4_wrong_size_words
    CORRECT_FONT_SIZE_H4 = 177800  # state the specified font size for Heading 4 style and store in the variable CORRECT_FONT_SIZE_H4; 177800/12700 = 14pt
    for paragraph in WordFile.paragraphs:
        if 'Heading 4' == paragraph.style.name:
            for run in paragraph.runs:
                # add fonts from each run into the set h4_sizes
                h4_size.add(run.font.size)
                # check if font sizes are unacceptable, if so, store in the set h4_wrong_size
                if run.font.size != CORRECT_FONT_SIZE_H4:
                    if run.font.size is not None:
                        h4_wrong_size.add(run.font.size / 12700)
                    else:
                        h4_wrong_size.add(run.font.size)
                    # append Heading 4 text that contain unacceptable font sizes in the sorted list h4_wrong_size_words
                    h4_wrong_size_words.append(run.text)

        # check if all elements in h4_size are not CORRECT_FONT_SIZE_H4 and if h4_size is not empty
    if {CORRECT_FONT_SIZE_H4} != h4_size and len(h4_size) != 0:
        # print this if all elements in h4_size are not CORRECT_FONT_SIZE_H4 and if h4_size is not empty and print h4_wrong_size and h4_wrong_size_words contents
        st.error(f'''
        ❌ Heading 4 style text have incorrect font sizes(s): {', '.join(map(str, h4_wrong_size))}  
        🡆 Incorrect font size(s) in Heading 4 style text found here: {' >> '.join(map(str, h4_wrong_size_words))}
        ''')
        # check if h4_size is empty, if so this means that Heading 4 style was not found
    elif len(h4_size) == 0:
        # print this if h4_size is empty, since Heading 4 style was not found
        st.info("ℹ️Heading 4 style text font size not found as this style was not used.")

    # TODO Font Size code for Normal style text
    # check font size for Normal style text (body text/paragraphs) ********************
    norm_size = set()  # store all Normal style font sizes in the set norm_size
    norm_wrong_size = set()  # store unacceptable Normal style font sizes in the set norm_wrong_size
    norm_wrong_size_words = sorted(set())  # store Normal text that are in unacceptable font size in the sorted list norm_wrong_size_words
    CORRECT_FONT_SIZE_NORM = 152400  # state the specified font size for Normal style and store in the variable CORRECT_FONT_SIZE_NORM; 152400/12700 = 12pt
    for paragraph in WordFile.paragraphs:
        if 'Normal' == paragraph.style.name:
            for run in paragraph.runs:
                # add fonts from each run into the set norm_size
                norm_size.add(run.font.size)
                # check if font sizes are unacceptable, if so, store in the set norm_wrong_size
                if run.font.size != CORRECT_FONT_SIZE_NORM:
                    if run.font.size is not None:
                        norm_wrong_size.add(run.font.size / 12700)
                    else:
                        norm_wrong_size.add(run.font.size)
                    # append Normal text that contain unacceptable font sizes in the sorted list norm_wrong_size_words
                    norm_wrong_size_words.append(run.text)

        # check if all elements in h1_font are not CORRECT_FONT_SIZE_NORM
    if {CORRECT_FONT_SIZE_NORM} != norm_size and len(norm_size) != 0:
        # print this if all elements in norm_font are not CORRECT_FONT_SIZE_NORM  and print norm_wrong_size and norm_wrong_size_words contents
        st.error(f'''
        ❌ Normal style text have incorrect font sizes(s): {', '.join(map(str, norm_wrong_size))}  
        🡆 Incorrect font size(s) in Normal style text found here: {' >> '.join(map(str, norm_wrong_size_words))}
        ''')
    elif len(norm_size) == 0:
        # print this if norm_size is empty, which means that Normal style was not found
        st.info("ℹ️Normal style text font size not found as this style was not used.")

    # TODO Font Size code for List Paragraph style text
    # check font size for List Paragraph style text (bullet list) ********************
    list_size = set()  # store all List Paragraph style font sizes in the set list_size
    list_wrong_size = set()  # store unacceptable List Paragraph style font sizes in the set list_wrong_size
    list_wrong_size_words = sorted(set())  # store List Paragraph text that are in unacceptable font size in the sorted list list_wrong_size_words
    CORRECT_FONT_SIZE_LIST = 152400  # state the specified font size for List Paragraph style and store in the variable CORRECT_FONT_SIZE_LIST; 152400/12700 = 12pt
    for paragraph in WordFile.paragraphs:
        if 'List Paragraph' == paragraph.style.name:
            for run in paragraph.runs:
                # add font sizes from each run into the set list_size
                list_size.add(run.font.size)
                # check if font sizes are unacceptable, if so, store in the set list_wrong_size
                if run.font.size != CORRECT_FONT_SIZE_LIST:
                    if run.font.size is not None:
                        list_wrong_size.add(run.font.size / 12700)
                    else:
                        list_wrong_size.add(run.font.size)
                    # append List Paragraph text that contain unacceptable font sizes in the sorted list list_wrong_size_words
                    list_wrong_size_words.append(run.text)

        # check if all elements in list_size are not CORRECT_FONT_SIZE_LIST and if list_size is not empty
    if {CORRECT_FONT_SIZE_LIST} != list_size and len(list_size) != 0:
        # print this if all elements in cap_size are not CORRECT_FONT_SIZE_LIST and if cap_size is not empty and print list_wrong_size and list_wrong_size_words contents
        st.error(f'''
        ❌ List Paragraph style text have incorrect font sizes(s): {', '.join(map(str, list_wrong_size))}  
        🡆 Incorrect font size(s) in List Paragraph style text found here: {' >> '.join(map(str, list_wrong_size_words))}
        ''')
        # check if list_size is empty, if so this means that List Paragraph style was not found
    elif len(list_size) == 0:
        # print this if list_size is empty, since List Paragraph style was not found
        st.info("ℹ️List Paragraph style text font size not found as this style was not used.")

    # TODO Font Size code for Caption style text
    # check font size for Caption style text (figure and table captions) ********************
    cap_size = set()  # store all Caption style font sizes in the set cap_size
    cap_wrong_size = set()  # store unacceptable Caption style font sizes in the set cap_wrong_size
    cap_wrong_size_words = sorted(set())  # store Caption text that are in unacceptable font size in the sorted list cap_wrong_size_words
    CORRECT_FONT_SIZE_CAP = 152400  # state the specified font size for Caption style and store in the variable CORRECT_FONT_SIZE_CAP; 152400/12700 = 12pt
    for paragraph in WordFile.paragraphs:
        if 'Caption' == paragraph.style.name:
            for run in paragraph.runs:
                # add font sizes from each run into the set cap_size
                cap_size.add(run.font.size)
                # check if font sizes are unacceptable, if so, store in the set cap_wrong_size
                if run.font.size != CORRECT_FONT_SIZE_CAP:
                    if run.font.size is not None:
                        cap_wrong_size.add(run.font.size / 12700)
                    else:
                        cap_wrong_size.add(run.font.size)
                    # append Caption text that contain unacceptable font sizes in the sorted list cap_wrong_size_words
                    cap_wrong_size_words.append(run.text)

        # check if all elements in cap_size are not CORRECT_FONT_SIZE_CAP and if cap_size is not empty
    if {CORRECT_FONT_SIZE_CAP} != cap_size and len(cap_size) != 0:
        # print this if all elements in cap_size are not CORRECT_FONT_SIZE_CAP and if cap_size is not empty and print cap_wrong_font and cap_wrong_font_words contents
        st.error(f'''
        ❌ Caption style text have incorrect font sizes(s): {', '.join(map(str, cap_wrong_size))}  
        🡆 Incorrect font size(s) in Caption style text found here: {' >> '.join(map(str, cap_wrong_size_words))}
        ''')
        # check if cap_size is empty, if so this means that Caption style was not found
    elif len(cap_size) == 0:
        # print this if cap_size is empty, since Caption style was not found
        st.info("ℹ️Caption style text font size not found as this style was not used.")


# font colour program function
def font_colour():
    # TODO start of Font Colour code --------------------
    # add font colour program banner
    st.subheader("**Font Colour**")
    # access Word document file
    WordFile = Document(uploaded_file)

    # TODO Font Colour code for Title style text
    # check font colour for Title style text (text on title page) ********************
    title_color = set()  # store all Title style font colors in the set title_color
    title_wrong_color = set()  # store unacceptable Title style font colors in the set title_wrong_color
    title_wrong_color_words = sorted(set())  # store Title text that are in unacceptable font color in the sorted list title_wrong_color_words
    CORRECT_FONT_COLOUR_TITLE = None  # state the specified font colour for Title style and store in the variable CORRECT_FONT_COLOUR_TITLE; None = default = Black (Automatic)
    for paragraph in WordFile.paragraphs:
        if 'Title' == paragraph.style.name:
            for run in paragraph.runs:
                # add font colors from each run into the set list_color
                title_color.add(run.font.color.rgb)
                # check if font colors are unacceptable, if so, store in the set title_wrong_color
                if run.font.color.rgb is not CORRECT_FONT_COLOUR_TITLE:
                    title_wrong_color.add(run.font.color.rgb)
                    # append Title text that contain unacceptable font colors in the sorted list title_wrong_color_words
                    title_wrong_color_words.append(run.text)

        # check if all elements in title_color are not CORRECT_FONT_COLOUR_TITLE and if title_color is not empty
    if title_color != {CORRECT_FONT_COLOUR_TITLE} and len(title_color) != 0:
        # print this if all elements in title_color are not CORRECT_FONT_COLOUR_TITLE and print title_wrong_color and title_wrong_words content
        st.error(f'''
        ❌ Title style text have incorrect font colour(s): {', '.join(map(str, title_wrong_color))}  
        🡆 Incorrect font colours(s) in Title style text found here: {' >> '.join(map(str, title_wrong_color_words))}
        ''')
        # check if title_color is empty, if so this means that Title style was not found
    elif len(title_color) == 0:
        # print this if title_color is empty, which means that Title style was not found
        st.info("ℹ️Title style text font colour not found as this style was not used.")

    # TODO Font Colour code for Heading 1 style text
    # check font colour for Heading 1 style text (main headings) ********************
    h1_color = set()  # store all Heading 1 style font colors in the set h1_color
    h1_wrong_color = set()  # store unacceptable Heading 1 style font colors in the set h1_wrong_color
    h1_wrong_color_words = sorted(set())  # store Heading 1 text that are in unacceptable font color in the sorted list h1_wrong_color_words
    CORRECT_FONT_COLOUR_H1 = None  # state the specified font colour for Title style and store in the variable CORRECT_FONT_COLOUR_H1; None = default = Black (Automatic)
    for paragraph in WordFile.paragraphs:
        if 'Heading 1' == paragraph.style.name:
            for run in paragraph.runs:
                # add font colors from each run into the set h1_color
                h1_color.add(run.font.color.rgb)
                # check if font colors are unacceptable, if so, store in the set h1_wrong_color
                if run.font.color.rgb is not CORRECT_FONT_COLOUR_H1:
                    h1_wrong_color.add(run.font.color.rgb)
                    # append Heading 1 text that contain unacceptable font colors in the sorted list h1_wrong_color_words
                    h1_wrong_color_words.append(run.text)

        # check if all elements in h1_color are not CORRECT_FONT_COLOUR_H1 and if h1_color is not empty
    if h1_color != {CORRECT_FONT_COLOUR_H1} and len(h1_color) != 0:
        # print this if all elements in h1_color are not CORRECT_FONT_COLOUR_H1 and print h1_wrong_color and h1_wrong_color_words content
        st.error(f'''
        ❌ Heading 1 style text have incorrect font colour(s): {', '.join(map(str, h1_wrong_color))}  
        🡆 Incorrect font colours(s) in Heading 1 style text found here: {' >> '.join(map(str, h1_wrong_color_words))}
        ''')
        # check if h1_color is empty, if so this means that Heading 1 style was not found
    elif len(h1_color) == 0:
        # print this if h1_color is empty, if so, this means that Heading 1 style was not found
        st.info("ℹ️Heading 1 style text font colour not found as this style was not used.")

    # TODO Font Colour code for Heading 2 style text
    # check font colour for Heading 2 style text (sub headings) ********************
    h2_color = set()  # store all Heading 2 style font colors in the set h2_color
    h2_wrong_color = set()  # store unacceptable Heading 2 style font colors in the set h2_wrong_color
    h2_wrong_color_words = sorted(set())  # store Heading 2 text that are in unacceptable font color in the sorted list h2_wrong_color_words
    CORRECT_FONT_COLOUR_H2 = None  # state the specified font colour for Title style and store in the variable CORRECT_FONT_COLOUR_H2 ; None = default = Black (Automatic)
    for paragraph in WordFile.paragraphs:
        if 'Heading 2' == paragraph.style.name:
            for run in paragraph.runs:
                # add font colors from each run into the set h1_color
                h2_color.add(run.font.color.rgb)
                # check if font colors are unacceptable, if so, store in the set h2_wrong_color
                if run.font.color.rgb is not CORRECT_FONT_COLOUR_H2:
                    h2_wrong_color.add(run.font.color.rgb)
                    # append Heading 2 text that contain unacceptable font colors in the sorted list h2_wrong_color_words
                    h2_wrong_color_words.append(run.text)

        # check if all elements in h2_color are not CORRECT_FONT_COLOUR_H2 and if h2_color is not empty
    if h2_color != {CORRECT_FONT_COLOUR_H2} and len(h2_color) != 0:
        # print this if all elements in h2_color are not CORRECT_FONT_COLOUR_H2 and print h2_wrong_color and h2_wrong_color_words content
        st.error(f'''
        ❌ Heading 2 style text have incorrect font colour(s): {', '.join(map(str, h2_wrong_color))}  
        🡆 Incorrect font colours(s) in Heading 2 style text found here: {' >> '.join(map(str, h2_wrong_color_words))}
        ''')
        # check if h2_color is empty, if so this means that Heading 2 style was not found
    elif len(h2_color) == 0:
        # print this if h2_color is empty, which means that Heading 2 style was not found
        st.info("ℹ️Heading 2 style text font colour not found as this style was not used.")

    # TODO Font Colour code for Heading 3 style text
    # check font colour for Heading 3 style text (sub headings) ********************
    h3_color = set()  # store all Heading 3 style font colors in the set h3_color
    h3_wrong_color = set()  # store unacceptable Heading 3 style font colors in the set h3_wrong_color
    h3_wrong_color_words = sorted(set())  # store Heading 3 text that are in unacceptable font color in the sorted list h3_wrong_color_words
    CORRECT_FONT_COLOUR_H3 = None  # state the specified font colour for Title style and store in the variable CORRECT_FONT_COLOUR_H3; None = default = Black (Automatic)
    for paragraph in WordFile.paragraphs:
        if 'Heading 3' == paragraph.style.name:
            for run in paragraph.runs:
                # add font colors from each run into the set h3_color
                h3_color.add(run.font.color.rgb)
                # check if font colors are unacceptable, if so, store in the set h3_wrong_color
                if run.font.color.rgb is not CORRECT_FONT_COLOUR_H3:
                    h3_wrong_color.add(run.font.color.rgb)
                    # append Heading 3 text that contain unacceptable font colors in the sorted list h3_wrong_color_words
                    h3_wrong_color_words.append(run.text)

        # check if all elements in h3_color are not CORRECT_FONT_COLOUR_H3 and if h3_color is not empty
    if h3_color != {CORRECT_FONT_COLOUR_H3} and len(h3_color) != 0:
        # print this if all elements in h3_color are not CORRECT_FONT_COLOUR_H3 and print h3_wrong_color and h3_wrong_color_words content
        st.error(f'''
        ❌ Heading 3 style text have incorrect font colour(s): {', '.join(map(str, h3_wrong_color))}  
        🡆 Incorrect font colours(s) in Heading 3 style text found here: {' >> '.join(map(str, h3_wrong_color_words))}
        ''')
        # check if h3_color is empty, if so this means that Heading 3 style was not found
    elif len(h3_color) == 0:
        # print this if h3_color is empty, which means that Heading 3 style was not found
        st.info("ℹ️Heading 3 style text font colour not found as this style was not used.")

    # TODO Font Colour code for Heading 4 style text
    # check font colour for Heading 4 style text (sub headings) ********************
    h4_color = set()  # store all Heading 4 style font colors in the set h4_color
    h4_wrong_color = set()  # store unacceptable Heading 4 style font colors in the set h4_wrong_color
    h4_wrong_color_words = sorted(set())  # store Heading 4 text that are in unacceptable font color in the sorted list h4_wrong_color_words
    CORRECT_FONT_COLOUR_H4 = None  # state the specified font colour for Title style and store in the variable CORRECT_FONT_COLOUR_H4; None = default = Black (Automatic)
    for paragraph in WordFile.paragraphs:
        if 'Heading 4' == paragraph.style.name:
            for run in paragraph.runs:
                # add font colors from each run into the set h4_color
                h4_color.add(run.font.color.rgb)
                # check if font colors are unacceptable, if so, store in the set h4_wrong_color
                if run.font.color.rgb is not CORRECT_FONT_COLOUR_H4:
                    h4_wrong_color.add(run.font.color.rgb)
                    # append Heading 4 text that contain unacceptable font colors in the sorted list h4_wrong_color_words
                    h4_wrong_color_words.append(run.text)

        # check if all elements in h4_color are not CORRECT_FONT_COLOUR_H4 and if h4_color is not empty
    if h4_color != {CORRECT_FONT_COLOUR_H4} and len(h4_color) != 0:
        # print this if all elements in h4_color are not CORRECT_FONT_COLOUR_H4 and print h4_wrong_color and h4_wrong_color_words content
        st.error(f'''
        ❌ Heading 4 style text have incorrect font colour(s): {', '.join(map(str, h4_wrong_color))}  
        🡆 Incorrect font colours(s) in Heading 4 style text found here: {' >> '.join(map(str, h4_wrong_color_words))}
        ''')
        # check if h4_color is empty, if so this means that Heading 4 style was not found
    elif len(h4_color) == 0:
        # print this if h4_color is empty, which means that Heading 4 style was not found
        st.info("ℹ️Heading 4 style text font colour not found as this style was not used.")

    # TODO Font Colour code for Normal style text
    # check font colour for normal style text (body text/paragraphs) ********************
    norm_color = set()  # store all Normal style font colors in the set norm_color
    norm_wrong_color = set()  # store unacceptable Normal style font colors in the set norm_wrong_color
    norm_wrong_color_words = sorted(set())  # store Normal text that are in unacceptable font color in the sorted list norm_wrong_color_words
    CORRECT_FONT_COLOUR_NORM = None  # state the specified font colour for Title style and store in the variable CORRECT_FONT_COLOUR_NORM; None = default = Black (Automatic)
    for paragraph in WordFile.paragraphs:
        if 'Normal' == paragraph.style.name:
            for run in paragraph.runs:
                # append font colors from each run into the set norm_color
                norm_color.add(run.font.color.rgb)
                # check if font colors are unacceptable, if so, store in the set norm_wrong_color
                if run.font.color.rgb is not CORRECT_FONT_COLOUR_NORM:
                    norm_wrong_color.add(run.font.color.rgb)
                    # store Normal text that contain unacceptable font colors in the sorted list norm_wrong_color_words
                    norm_wrong_color_words.append(run.text)

        # check if all elements in norm_color are not CORRECT_FONT_COLOUR_NORM
    if norm_color != {CORRECT_FONT_COLOUR_NORM} and len(norm_color) != 0:
        # print this if all elements in norm_color are not CORRECT_FONT_COLOUR_NORM and print norm_wrong_color and norm_wrong_color_words content
        st.error(f'''
        ❌ Normal style text have incorrect font colour(s): {', '.join(map(str, norm_wrong_color))}  
        🡆 Incorrect font colours(s) in Normal style text found here: {' >> '.join(map(str, norm_wrong_color_words))}
        ''')
    elif len(norm_color) == 0:
        # print this if norm_color is empty, which means that Normal style was not found
        st.info("ℹ️Normal style text font colour not found as this style was not used.")

    # TODO Font Colour code for List Paragraph style text
    # check font colour for List Paragraph style text (bullet list) ********************
    list_color = set()  # store all List Paragraph style font colors in the set list_color
    list_wrong_color = set()  # store unacceptable List Paragraph style font colors in the set list_wrong_color
    list_wrong_color_words = sorted(set())  # store List Paragraph text that are in unacceptable font color in the sorted list list_wrong_color_words
    CORRECT_FONT_COLOUR_LIST = None  # state the specified font colour for Title style and store in the variable CORRECT_FONT_COLOUR_LIST; None = default = Black (Automatic)
    for paragraph in WordFile.paragraphs:
        if 'List Paragraph' == paragraph.style.name:
            for run in paragraph.runs:
                # add font colors from each run into the set list_color
                list_color.add(run.font.color.rgb)
                # check if font colors are unacceptable, if so, store in the set list_wrong_color
                if run.font.color.rgb is not CORRECT_FONT_COLOUR_LIST:
                    list_wrong_color.add(run.font.color.rgb)
                    # append List Paragraph text that contain unacceptable font colors in the sorted list list_wrong_color_words
                    list_wrong_color_words.append(run.text)

        # check if all elements in list_color are not CORRECT_FONT_COLOUR_LIST and if list_color is not empty
    if list_color != {CORRECT_FONT_COLOUR_LIST} and len(list_color) != 0:
        # print this if all elements in list_color are not CORRECT_FONT_COLOUR_LIST and print list_wrong_color and list_wrong_color_words content
        st.error(f'''
        ❌ List Paragraph style text have incorrect font colour(s): {', '.join(map(str, list_wrong_color))}  
        🡆 Incorrect font colours(s) in List Paragraph style text found here: {' >> '.join(map(str, list_wrong_color_words))}
        ''')
        # check if list_color is empty, if so this means that List Paragraph style was not found
    elif len(list_color) == 0:
        # print this if list_color is empty, which means that List Paragraph style was not found
        st.info("ℹ️List Paragraph style text font colour not found as this style was not used.")

    # TODO Font Colour code for Caption style text
    # check font colour for Caption style text (figure and table captions) ********************
    cap_color = set()  # store all Caption style font colors in the set cap_color
    cap_wrong_color = set()  # store unacceptable Caption style font colors in the set cap_wrong_color
    cap_wrong_color_words = sorted(set())  # store Caption text that are in unacceptable font color in the sorted list cap_wrong_color_words
    CORRECT_FONT_COLOUR_CAP = None  # state the specified font colour for Title style and store in the variable CORRECT_FONT_COLOUR_CAP; None = default = Black (Automatic)
    for paragraph in WordFile.paragraphs:
        if 'Caption' == paragraph.style.name:
            for run in paragraph.runs:
                # add font colors from each run into the set cap_color
                cap_color.add(run.font.color.rgb)
                # check if font colors are unacceptable, if so, store in the set cap_wrong_color
                if run.font.color.rgb is not CORRECT_FONT_COLOUR_CAP:
                    cap_wrong_color.add(run.font.color.rgb)
                    # append Caption text that contain unacceptable font colors in the sorted list cap_wrong_color_words
                    cap_wrong_color_words.append(run.text)

        # check if all elements in cap_color are not CORRECT_FONT_COLOUR_CAP and if cap_color is not empty
    if cap_color != {CORRECT_FONT_COLOUR_CAP} and len(cap_color) != 0:
        # print this if all elements in cap_color are not CORRECT_FONT_COLOUR_CAP and print cap_wrong_color and cap_wrong_color_words content
        st.error(f'''
        ❌ Caption style text have incorrect font colour(s): {', '.join(map(str, cap_wrong_color))}  
        🡆 Incorrect font colours(s) in Caption style text found here: {' >> '.join(map(str, cap_wrong_color_words))}
        ''')
        # check if cap_color is empty, if so this means that Caption style was not found
    elif len(cap_color) == 0:
        # print this if cap_color is empty, which means that Caption style was not found
        st.info("ℹ️Caption style text font colour not found as this style was not used.")


# character formatting program function
def char_form():
    # TODO start of Character Formatting code --------------------
    # add character formatting program banner
    st.subheader('**Character Formatting**')
    # access Word document file
    WordFile = Document(uploaded_file)

    # TODO Character Formatting code for Heading 1 style text
    # Check all Heading 1 text and determine if it is in bold and or italics ********************
    h1_bold = set()  # store True and None for instances where Heading 1 is and not bold (None == off, True == on)
    h1_italic = set()  # store True and None for instances where Heading 1 is and not italicised (None == off, True == on)
    h1_not_bold_words = sorted(set())  # store Heading 1 text that are not in bold in the sorted list h1_not_bold_words
    h1_italic_words = sorted(set())  # store Heading 1 text that are italicised in the sorted list h1_italic_words
    h1_bold_italic_words = sorted(set())  # store Heading 1 text that are bold and italicised in the sorted list h1_bold_italic_words
    for paragraph in WordFile.paragraphs:
        if 'Heading 1' == paragraph.style.name:
            for run in paragraph.runs:
                # append Heading 1 bold status from each run into the set h1_bold
                h1_bold.add(run.font.bold)
                # append Heading 1 italic status from each run into the set h1_italic
                h1_italic.add(run.font.italic)
                # check if Heading 1 is not bold
                if run.font.bold is None:
                    # append Heading 1 text that is not bold in the sorted list h1_not_bold_words
                    h1_not_bold_words.append(run.text)
                # check if Heading 1 is italicised
                if run.font.italic is True:
                    # append Heading 1 text that is italicised in the sorted list h1_italic_words
                    h1_italic_words.append(run.text)
                # check if Heading 1 is bold and italicised
                if run.font.bold is True and run.font.italic is True:
                    # append Heading 1 text that is bold and italicised in the sorted list h1_bold_italic_words
                    h1_bold_italic_words.append(run.text)

        # check if all elements in h1_bold are None and if all elements in h1_italic are None
    if all(x is None for x in h1_bold) and all(x is None for x in h1_italic) and len(h1_bold) != 0 and len(h1_italic) != 0:
        st.error("❌ All Heading 1 style text are not bold.")
        # check if any of the elements of h1_bold are None
    elif any(x is None for x in h1_bold) and len(h1_bold) != 0 and len(h1_italic) != 0:
        st.error(f"❌ Non-bold Heading 1 style text found: {' >> '.join(map(str, h1_not_bold_words))}")
        # check if all elements of h1_italic are True and if all elements of h1_bold are None

    if all(x is True for x in h1_italic) and all(x is None for x in h1_bold) and len(h1_bold) != 0 and len(h1_italic) != 0:
        st.error("❌ All Heading 1 style text are italicised.")
        # check if any of the elements of h1_italic are True
    elif any(x is True for x in h1_italic) and len(h1_bold) != 0 and len(h1_italic) != 0:
        st.error(f"❌ Italicised Heading 1 style text found: {' >> '.join(map(str, h1_italic_words))}")
        # check if all elements in h1_bold and h1_italic are True

    if all(x is True for x in h1_bold) and all(x is True for x in h1_italic) and len(h1_bold) != 0 and len(h1_italic) != 0:
        st.error("❌ All Heading 1 style text are bold and italicised.")
        # check if any of the elements of h1_bold and h1_italic are True
    elif any(x is True for x in h1_bold) and any(x is True for x in h1_italic) and len(h1_bold) != 0 and len(h1_italic) != 0:
        st.error(f"❌ Bold and italicised Heading 1 style text found: {' >> '.join(map(str, h1_bold_italic_words))}")
        # check if both h1_bold and h1_italic are empty

    if len(h1_bold) == 0 and len(h1_italic) == 0:
        st.info("ℹ️Heading 1 style text bold and italic check not conducted as this style was not used.")

    # TODO Character Formatting code for Heading 2 style text
    # Check all Heading 2 text and determine if it is in bold and or italics ********************
    h2_bold = set()  # store True and None for instances where Heading 2 is and not bold (None == off, True == on)
    h2_italic = set()  # store True and None for instances where Heading 2 is and not italicised (None == off, True == on)
    h2_not_italic_words = sorted(set())  # store Heading 2 text that are not italicised in the sorted list h2_not_italic_words
    h2_bold_words = sorted(set())  # store Heading 2 text that are bold in the sorted list h2_bold_words
    h2_bold_italic_words = sorted(set())  # store Heading 2 text that are bold and italicised in the sorted list h2_bold_italic_words
    for paragraph in WordFile.paragraphs:
        if 'Heading 2' == paragraph.style.name:
            for run in paragraph.runs:
                # add Heading 2 bold status from each run into the set h2_bold
                h2_bold.add(run.font.bold)
                # add Heading 2 italic status from each run into the set h2_italic
                h2_italic.add(run.font.italic)
                # check if Heading 2 is not italicised
                if run.font.italic is None:
                    # append Heading 2 text that is not italicised in the sorted list h2_not_italic_words
                    h2_not_italic_words.append(run.text)
                # check if Heading 2 is bold
                if run.font.bold is True:
                    # append Heading 2 text that is bold in the sorted list h2_bold_words
                    h2_bold_words.append(run.text)
                # check if Heading 2 is bold and italicised
                if run.font.bold is True and run.font.italic is True:
                    # append Heading 2 text that is bold and italicised in the sorted list h2_bold_italic_words
                    h2_bold_italic_words.append(run.text)

        # check if all elements in h2_italic are None and if all elements in h2_bold are None
    if all(x is None for x in h2_italic) and all(x is None for x in h2_bold) and len(h2_bold) != 0 and len(h2_italic) != 0:
        st.error("❌ All Heading 2 style text are not italicised.")
        # check if any elements of h2_italic are None
    elif any(x is None for x in h2_italic) and len(h2_bold) != 0 and len(h2_italic) != 0:
        st.error(f"❌ Non-italicised Heading 2 style text found: {' >> '.join(map(str, h2_not_italic_words))}")
        # check if all elements in h2_bold are True and if all elements in h2_italic are None

    if all(x is True for x in h2_bold) and all(x is None for x in h2_italic) and len(h2_bold) != 0 and len(h2_italic) != 0:
        st.error("❌ All Heading 2 style text are bold.")
        # check if any of the elements of h2_bold are True
    elif any(x is True for x in h2_bold) and len(h2_bold) != 0 and len(h2_italic) != 0:
        st.error(f"❌ Bold Heading 2 style text found: {' >> '.join(map(str, h2_bold_words))}")
        # check if all elements in h2_bold and h2_italic are True

    if all(x is True for x in h2_bold) and all(x is True for x in h2_italic) and len(h2_bold) != 0 and len(h2_italic) != 0:
        st.error("❌ All Heading 2 style text are bold and italicised.")
        # check if any of the elements of h2_bold and h2_italic are True
    elif any(x is True for x in h2_bold) and any(x is True for x in h2_italic) and len(h2_bold) != 0 and len(h2_italic) != 0:
        st.error(f"❌ Bold and italicised Heading 2 style text found: {' >> '.join(map(str, h2_bold_italic_words))}")
        # check if both h2_bold and h2_italic are empty

    if len(h2_bold) == 0 and len(h2_italic) == 0:
        st.info("ℹ️Heading 2 style text bold and italic check not conducted as this style was not used.")

    # TODO Character Formatting code for Heading 3 style text
    # Check all Heading 3 text and determine if it is in bold and or italics ********************
    h3_bold = set()  # store True and None for instances where Heading 3 is and not bold (None == off, True == on)
    h3_italic = set()  # store True and None for instances where Heading 3 is and not italicised (None == off, True == on)
    h3_not_italic_words = sorted(set())  # store Heading 3 text that are not italicised in the sorted list h3_not_italic_words
    h3_bold_words = sorted(set())  # store Heading 3 text that are bold in the sorted list h3_bold_words
    h3_bold_italic_words = sorted(set())  # store Heading 3 text that are bold and italicised in the sorted list h3_bold_italic_words
    for paragraph in WordFile.paragraphs:
        if 'Heading 3' == paragraph.style.name:
            for run in paragraph.runs:
                # add Heading 3 bold status from each run into the set h3_bold
                h3_bold.add(run.font.bold)
                # add Heading 3 italic status from each run into the set h3_italic
                h3_italic.add(run.font.italic)
                # check if Heading 3 is not italicised
                if run.font.italic is None:
                    # append Heading 3 text that is not italicised in the sorted list h3_not_italic_words
                    h3_not_italic_words.append(run.text)
                # check if Heading 3 is bold
                if run.font.bold is True:
                    # append Heading 3 text that is bold in the sorted list h3_bold_words
                    h3_bold_words.append(run.text)
                # check if Heading 3 is bold and italicised
                if run.font.bold is True and run.font.italic is True:
                    # append Heading 3 text that is bold and italicised in the sorted list h3_bold_italic_words
                    h3_bold_italic_words.append(run.text)

        # check if all elements in h3_italic are None and if all elements in h3_bold are None
    if all(x is None for x in h3_italic) and all(x is None for x in h3_bold) and len(h3_bold) != 0 and len(h3_italic) != 0:
        st.error("❌ All Heading 3 style text are not italicised.")
        # check if any elements of h3_italic are None
    elif any(x is None for x in h3_italic) and len(h3_bold) != 0 and len(h3_italic) != 0:
        st.error(f"❌ Non-italicised Heading 3 style text found: {' >> '.join(map(str, h3_not_italic_words))}")
        # check if all elements in h3_bold are True and if all elements in h3_italic are None

    if all(x is True for x in h3_bold) and all(x is None for x in h3_italic) and len(h3_bold) != 0 and len(h3_italic) != 0:
        st.error("❌ All Heading 3 style text are bold.")
        # check if any of the elements of h3_bold are True
    elif any(x is True for x in h3_bold) and len(h3_bold) != 0 and len(h3_italic) != 0:
        st.error(f"❌ Bold Heading 3 style text found: {' >> '.join(map(str, h3_bold_words))}")
        # check if all elements in h3_bold and h3_italic are True

    if all(x is True for x in h3_bold) and all(x is True for x in h3_italic) and len(h3_bold) != 0 and len(h3_italic) != 0:
        st.error("❌ All Heading 3 style text are bold and italicised.")
        # check if any of the elements of h3_bold and h3_italic are True
    elif any(x is True for x in h3_bold) and any(x is True for x in h3_italic) and len(h3_bold) != 0 and len(h3_italic) != 0:
        st.error(f"❌ Bold and italicised Heading 3 style text found: {' >> '.join(map(str, h3_bold_italic_words))}")
        # check if both h3_bold and h3_italic are empty

    if len(h3_bold) == 0 and len(h3_italic) == 0:
        st.info("ℹ️Heading 3 style text bold and italic check not conducted as this style was not used.")

    # TODO Character Formatting code for Heading 4 style text
    # Check all Heading 4 text and determine if it is in bold and or italics ********************
    h4_bold = set()  # store True and None for instances where Heading 4 is and not bold (None == off, True == on)
    h4_italic = set()  # store None and False for instances where Heading 4 is and not italicised (False == off, None == on)
    h4_not_italic_words = sorted(set())  # store Heading 4 text that are not italicised in the sorted list h4_not_italic_words
    h4_bold_words = sorted(set())  # store Heading 4 text that are bold in the sorted list h4_bold_words
    h4_bold_italic_words = sorted(set())  # store Heading 4 text that are bold and italicised in the sorted list h4_bold_italic_words
    for paragraph in WordFile.paragraphs:
        if 'Heading 4' == paragraph.style.name:
            for run in paragraph.runs:
                # add Heading 4 bold status from each run into the set h4_bold
                h4_bold.add(run.font.bold)
                # add Heading 4 italic status from each run into the set h4_italic
                h4_italic.add(run.font.italic)
                # check if Heading 4 is not italicised
                if run.font.italic is False:
                    # append Heading 4 text that is not italicised in the sorted list h4_not_italic_words
                    h4_not_italic_words.append(run.text)
                # check if Heading 4 is bold
                if run.font.bold is True:
                    # append Heading 4 text that is bold in the sorted list h4_bold_words
                    h4_bold_words.append(run.text)
                # check if Heading 4 is bold and italicised
                if run.font.bold is True and run.font.italic is None:
                    # append Heading 4 text that is bold and italicised in the sorted list h4_bold_italic_words
                    h4_bold_italic_words.append(run.text)

        # check if all elements in h4_italic are False and if all elements in h4_bold are None
    if all(x is False for x in h4_italic) and all(x is None for x in h4_bold) and len(h4_bold) != 0 and len(h4_italic) != 0:
        st.error("❌ All Heading 4 style text are not italicised.")
        # check if any elements of h4_italic are False
    elif any(x is False for x in h4_italic) and len(h4_bold) != 0 and len(h4_italic) != 0:
        st.error(f"❌ Non-italicised Heading 4 style text found: {' >> '.join(map(str, h4_not_italic_words))}")
        # check if all elements in h4_bold are True and if all elements in h4_italic are False

    if all(x is True for x in h4_bold) and all(x is False for x in h4_italic) and len(h4_bold) != 0 and len(h4_italic) != 0:
        st.error("❌ All Heading 4 style text are bold.")
        # check if any of the elements of h4_bold are True
    elif any(x is True for x in h4_bold) and len(h4_bold) != 0 and len(h4_italic) != 0:
        st.error(f"❌ Bold Heading 4 style text found: {' >> '.join(map(str, h4_bold_words))}")
        # check if all elements in h4_bold and h4_italic are True and None, respectively

    if all(x is True for x in h4_bold) and all(x is None for x in h4_italic) and len(h4_bold) != 0 and len(h4_italic) != 0:
        st.error("❌ All Heading 4 style text are bold and italicised.")
        # check if any of the elements in h4_bold and h4_italic are True and None, respectively
    elif any(x is True for x in h4_bold) and any(x is None for x in h4_italic) and len(h4_bold) != 0 and len(h4_italic) != 0:
        st.error(f"❌ Bold and italicised Heading 4 style text found: {' >> '.join(map(str, h4_bold_italic_words))}")
        # check if both h4_bold and h4_italic are empty

    if len(h4_bold) == 0 and len(h4_italic) == 0:
        st.info("ℹ️Heading 4 style text bold and italic check not conducted as this style was not used.")


# paragraph alignment program function
def para_align():
    # TODO start of Paragraph Alignment code --------------------
    # add paragraph alignment program banner
    st.subheader("**Paragraph Alignment**")
    # access Word document file
    WordFile = Document(uploaded_file)

    # TODO Paragraph Alignment code for Normal style text
    # check the paragraph alignment of all text formatted with the 'Normal' style (body text\paragraphs) ********************
    norm_align = set()  # store all paragraph alignments in the set norm_align
    norm_wrong_align = set()  # store unacceptable paragraph alignments in the set norm_wrong_align
    norm_wrong_align_words = sorted(set())  # store text from unacceptable paragraph alignments in the sorted list norm_wrong_align
    CORRECT_ALIGN_NORM = 3  # state the specified alignment for Normal style and store in the variable CORRECT_ALIGN_NORM; None = Left, 1 = Center, 2 = Right, 3 = Justify
    for paragraph in WordFile.paragraphs:
        if 'Normal' == paragraph.style.name:
            for run in paragraph.runs:
                # add paragraph alignment from each run into the set norm_align
                norm_align.add(paragraph.alignment)
                # check if paragraph alignment is unacceptable, if so, store in the set norm_wrong_align
                if paragraph.alignment != CORRECT_ALIGN_NORM:
                    norm_wrong_align.add(paragraph.alignment)
                    # append Normal text that contain unacceptable paragraph alignment in the sorted list norm_wrong_align_words
                    norm_wrong_align_words.append(paragraph.text)

        # check if all elements in norm_align are not CORRECT_ALIGN_NORM
    if {CORRECT_ALIGN_NORM} != norm_align and len(norm_align) != 0:
        # print this if all elements in norm_align are not CORRECT_ALIGN_NORM and print norm_wrong_align and norm_wrong_align_words contents
        st.error(f'''
        ❌ Normal style text has incorrect paragraph alignment: {', '.join(map(str, norm_wrong_align))}  
        🡆 Incorrect paragraph alignment for Normal style text found here: {' >> '.join(map(str, norm_wrong_align_words))}
        ''')
    # check if the set norm_align is empty
    elif len(norm_align) == 0:
        st.info("ℹ️Normal style text paragraph alignment not found as this style was not used.")

    # TODO Paragraph Alignment code for List Paragraph style text
    # check the paragraph alignment of all text formatted with the 'List Paragraph' style (bullet lists) ********************
    list_align = set()  # store all paragraph alignments in the set list_align
    list_wrong_align = set()  # store unacceptable paragraph alignments in the set list_wrong_align
    list_wrong_align_words = sorted(set())  # store text from unacceptable bullet list alignments in the sorted list list_wrong_align
    CORRECT_ALIGN_LIST = 3  # state the specified alignment for List Paragraph style and store in the variable CORRECT_ALIGN_LIST; None = Left, 1 = Center, 2 = Right, 3 = Justify
    for paragraph in WordFile.paragraphs:
        if 'List Paragraph' == paragraph.style.name:
            for run in paragraph.runs:
                # add paragraph alignment from each run into the set list_align
                list_align.add(paragraph.alignment)
                # check if paragraph alignment is unacceptable, if so, store in the set list_wrong_align
                if paragraph.alignment != CORRECT_ALIGN_LIST:
                    list_wrong_align.add(paragraph.alignment)
                    # append list paragraph text that contain unacceptable paragraph alignment in the sorted list list_wrong_align_words
                    list_wrong_align_words.append(paragraph.text)

        # check if all elements in list_align are not CORRECT_ALIGN_LIST and if the list list_align is not empty
    if {CORRECT_ALIGN_LIST} != list_align and len(list_align) != 0:
        # print this if all elements in list_align are not CORRECT_ALIGN_LIST and print list_wrong_align and list_wrong_align_words contents
        st.error(f'''
        ❌ List Paragraph style text has incorrect paragraph alignment: {', '.join(map(str, list_wrong_align))}  
        🡆 Incorrect paragraph alignment for List Paragraph style text found here: {' >> '.join(map(str, list_wrong_align_words))}
        ''')
        # check if the set list_align is empty
    elif len(list_align) == 0:
        st.info("ℹ️List Paragraph style text paragraph alignment not found as this style was not used.")


# spacing program function
def spacing():
    # TODO start of Spacing code --------------------
    # add spacing program banner
    st.subheader("**Spacing**")
    # access Word document file
    WordFile = Document(uploaded_file)

    # TODO spacing before code for Normal style text
    # Check spacing before paragraph for all text formatted in the 'Normal' style (body text/paragraph) ********************
    norm_para_b = set()  # store all spacing before paragraph values for Normal style text in the set norm_para_b
    norm_wrong_para_b = set()  # store unacceptable spacing before paragraph values for Normal style text in the set norm_wrong_para_b
    norm_wrong_para_b_words = sorted(set())  # store Normal style text that contain unacceptable spacing before paragraph values in the sorted set norm_wrong_para_b_words
    CORRECT_SPACE_BEFORE_NORM = None  # state the specified spacing before for Normal style and store in the variable CORRECT_SPACE_BEFORE_NORM; None = default = 0pt
    for paragraph in WordFile.paragraphs:
        if 'Normal' == paragraph.style.name:
            for run in paragraph.runs:
                # add Normal style text spacing before paragraph values in the set norm_para_b
                norm_para_b.add(paragraph.paragraph_format.space_before)
                # check if spacing before paragraph is unacceptable, if so, store in the set norm_wrong_para_b
                if paragraph.paragraph_format.space_before is not CORRECT_SPACE_BEFORE_NORM:
                    if paragraph.paragraph_format.space_before is not None:
                        norm_wrong_para_b.add(paragraph.paragraph_format.space_before / 12700)
                    else:
                        norm_wrong_para_b.add(paragraph.paragraph_format.space_before)
                    # append Normal style text that contain unacceptable spacing before paragraph values in the sorted list norm_wrong_para_b_words
                    norm_wrong_para_b_words.append(run.text)

        # check if all elements in norm_para_b are not CORRECT_SPACE_BEFORE_NORM
    if norm_para_b != {CORRECT_SPACE_BEFORE_NORM} and len(norm_para_b) != 0:
        # print this if all elements in norm_para_b are not CORRECT_SPACE_BEFORE_NORM and print norm_wrong_para_b content
        st.error(f'''
        ❌ Normal style text has incorrect spacing before paragraph: {', '.join(map(str, norm_wrong_para_b))}  
        🡆 Incorrect spacing before paragraph for Normal style text found here: {' >> '.join(map(str, norm_wrong_para_b_words))}
        ''')
    # check if norm_para_b is empty, if so this means that Normal style was not found
    elif len(norm_para_b) == 0:
        # print this if norm_para_b is empty, since Normal style was not found
        st.info("ℹ️Normal style text spacing before paragraph not found as this style was not used.")

    # TODO line spacing code for Normal style text
    # Check line spacing for all text formatted in the 'Normal' style (body text/paragraphs) ***************************
    norm_para_l = set()  # store all Normal text line spacing values in the set norm_para_l
    norm_wrong_para_l = set()  # store unacceptable Normal text line spacing values in the set norm_wrong_para_l
    norm_wrong_para_l_words = sorted(set())  # store Normal text with unacceptable line spacing in sorted list norm_wrong_para_l_words
    CORRECT_LINE_SPACE_NORM = 2.0  # state the specified line spacing for Normal style and store in the variable CORRECT_LINE_SPACE_NORM
    for paragraph in WordFile.paragraphs:
        if 'Normal' == paragraph.style.name:
            for run in paragraph.runs:
                # add Normal text line spacing values in the set norm_para_l
                norm_para_l.add(paragraph.paragraph_format.line_spacing)
                # check if Normal text line spacing is unacceptable, if so, store in the set norm_wrong_para_l
                if paragraph.paragraph_format.line_spacing != CORRECT_LINE_SPACE_NORM:
                    norm_wrong_para_l.add(paragraph.paragraph_format.line_spacing)
                    # append Normal text that contain unacceptable line spacing in the sorted list norm_wrong_para_l_words
                    norm_wrong_para_l_words.append(run.text)

        # check if all elements in norm_para_l are not CORRECT_LINE_SPACE_NORM
    if {CORRECT_LINE_SPACE_NORM} != norm_para_l and len(norm_para_l) != 0:
        # print this if all elements in norm_para_l are not CORRECT_LINE_SPACE_NORM and print norm_wrong_para_l and norm_wrong_para_l_words content
        st.error(f'''
        ❌ Normal style text has incorrect line spacing: {', '.join(map(str, norm_wrong_para_l))}  
        🡆 Incorrect line spacing in Normal style text found here: {' >> '.join(map(str, norm_wrong_para_l_words))}
        ''')
        # check if norm_para_l is empty, if so this means that Normal style was not found
    elif len(norm_para_l) == 0:
        # print this if norm_para_l is empty, since Normal style was not found
        st.info("ℹ️Normal style text line spacing not found as this style was not used.")

    # TODO spacing after code for Normal style text
    # Check spacing after paragraph for all text formatted in the 'Normal' style (body text/paragraph) ********************
    norm_para_a = set()  # store all spacing after paragraph values for Normal style text in the set norm_para_a
    norm_wrong_para_a = set()  # store unacceptable spacing after paragraph values for Normal style text in the set norm_wrong_para_a
    norm_wrong_para_a_words = sorted(set())  # store Normal style text that contain unacceptable spacing after paragraph values in the sorted set norm_wrong_para_a_words
    CORRECT_SPACE_AFTER_NORM = None  # state the specified spacing after for Normal style and store in the variable CORRECT_SPACE_AFTER_NORM; None = default = 8pt
    for paragraph in WordFile.paragraphs:
        if 'Normal' == paragraph.style.name:
            for run in paragraph.runs:
                # add Normal style text spacing after paragraph values in the set norm_para_a
                norm_para_a.add(paragraph.paragraph_format.space_after)
                # check if spacing after paragraph is unacceptable, if so, store in the set norm_wrong_para_a
                if paragraph.paragraph_format.space_after is not CORRECT_SPACE_AFTER_NORM:
                    if paragraph.paragraph_format.space_after is not None:
                        norm_wrong_para_a.add(paragraph.paragraph_format.space_after / 12700)
                    else:
                        norm_wrong_para_a.add(paragraph.paragraph_format.space_after)
                    # append Normal style text that contain unacceptable spacing after paragraph values in the sorted list norm_wrong_para_a_words
                    norm_wrong_para_a_words.append(run.text)

        # check if all elements in norm_para_a are not CORRECT_SPACE_AFTER_NORM
    if norm_para_a != {CORRECT_SPACE_AFTER_NORM} and len(norm_para_a) != 0:
        # print this if all elements in norm_para_a are not CORRECT_SPACE_AFTER_NORM and print norm_wrong_para_a and norm_wrong_para_a_words content
        st.error(f'''
        ❌ Normal style text has incorrect spacing after paragraph: {', '.join(map(str, norm_wrong_para_a))}  
        🡆 Incorrect spacing after paragraph for Normal style text found here: {' >> '.join(map(str, norm_wrong_para_a_words))}
        ''')
        # check if norm_para_a is empty, if so this means that Normal style was not found
    elif len(norm_para_a) == 0:
        # print this if norm_para_a is empty, since Normal style was not found
        st.info("ℹ️Normal style text spacing after paragraph not found as this style was not used.")

    # TODO spacing before code for List Paragraph style text
    # Check spacing before paragraph for all text formatted in the 'List Paragraph' style (bullet lists) ********************
    list_para_b = set()  # store all spacing before paragraph values for List Paragraph style text in the set list_para_b
    list_wrong_para_b = set()  # store unacceptable spacing before paragraph values for List Paragraph style text in the set list_wrong_para_b
    list_wrong_para_b_words = sorted(set())  # store List Paragraph style text that contain unacceptable spacing before paragraph values in the sorted list list_wrong_para_b_words
    CORRECT_SPACE_BEFORE_LIST = None  # state the specified spacing before for List Paragraph style and store in the variable CORRECT_SPACE_BEFORE_LIST; None = default = 0pt
    for paragraph in WordFile.paragraphs:
        if 'List Paragraph' == paragraph.style.name:
            for run in paragraph.runs:
                # add List Paragraph style text spacing before paragraph values in the set list_para_b
                list_para_b.add(paragraph.paragraph_format.space_before)
                # check if spacing before paragraph is unacceptable, if so, store in the set list_wrong_para_b
                if paragraph.paragraph_format.space_before is not CORRECT_SPACE_BEFORE_LIST:
                    if paragraph.paragraph_format.space_before is not None:
                        list_wrong_para_b.add(paragraph.paragraph_format.space_before / 12700)
                    else:
                        list_wrong_para_b.add(paragraph.paragraph_format.space_before)
                    # append List Paragraph style text that contain unacceptable spacing before paragraph values in the sorted list list_wrong_para_b_words
                    list_wrong_para_b_words.append(run.text)

        # check if all elements in list_para_b are not CORRECT_SPACE_BEFORE_LIST
    if list_para_b != {CORRECT_SPACE_BEFORE_LIST} and len(list_para_b) != 0:
        # print this if all elements in list_para_b are not CORRECT_SPACE_BEFORE_LIST and print list_wrong_para_b and list_wrong_para_b_words content
        st.error(f'''
        ❌ List Paragraph style text has incorrect spacing before paragraph: {', '.join(map(str, list_wrong_para_b))}  
        🡆 Incorrect spacing before paragraph for List Paragraph style text found here: {' >> '.join(map(str, list_wrong_para_b_words))}
        ''')
        # check if list_para_b is empty, if so this means that List Paragraph style was not found
    elif len(list_para_b) == 0:
        # print this if list_para_b is empty, since List Paragraph style was not found
        st.info("ℹ️List Paragraph style text spacing before paragraph not found as this style was not used.")

    # TODO line spacing code for List Paragraph style text
    # Check line spacing for all text formatted in the List Paragraph style (bullet list) ***************************
    list_para_l = set()  # store all List Paragraph style text line spacing values in the set list_para_l
    list_wrong_para_l = set()  # store unacceptable List Paragraph style text line spacing values in the set list_wrong_para_l
    list_wrong_para_l_words = sorted(set())  # store List Paragraph style text with unacceptable line spacing in the sorted list list_wrong_para_l_words
    CORRECT_LINE_SPACE_LIST = 1.0  # state the specified line spacing for List Paragraph style and store in the variable CORRECT_LINE_SPACE_LIST
    for paragraph in WordFile.paragraphs:
        if 'List Paragraph' == paragraph.style.name:
            for run in paragraph.runs:
                # add List Paragraph style text line spacing values in the set list_para_l
                list_para_l.add(paragraph.paragraph_format.line_spacing)
                # check if line spacing are unacceptable, if so, store in the set list_wrong_para_l
                if paragraph.paragraph_format.line_spacing != CORRECT_LINE_SPACE_LIST:
                    list_wrong_para_l.add(paragraph.paragraph_format.line_spacing)
                    # append List Paragraph style text that contain unacceptable line spacing in the sorted list list_wrong_para_l_words
                    list_wrong_para_l_words.append(run.text)

        # check if all elements in list_para_l are not CORRECT_LINE_SPACE_LIST
    if {CORRECT_LINE_SPACE_LIST} != list_para_l and len(list_para_l) != 0:
        # print this if all elements in list_para_l are not CORRECT_LINE_SPACE_LIST and print list_wrong_para_l and list_wrong_para_l_words content
        st.error(f'''
        ❌ List Paragraph style text has incorrect line spacing: {', '.join(map(str, list_wrong_para_l))}  
        🡆 Incorrect line spacing in List Paragraph style text found here: {' >> '.join(map(str, list_wrong_para_l_words))}
        ''')
        # check if list_para_l is empty, if so this means that List Paragraph style was not found
    elif len(list_para_l) == 0:
        # print this if list_para_l is empty, since List Paragraph style was not found
        st.info("ℹ️List Paragraph style text line spacing not found as this style was not used.")

    # TODO spacing after code for List Paragraph style text
    # Check spacing after paragraph for all text formatted in the 'List Paragraph' style (bullet list) ********************
    list_para_a = set()  # store all spacing after paragraph values for List Paragraph style text in the set list_para_a
    list_wrong_para_a = set()  # store unacceptable spacing after values for List Paragraph style text in the set list_wrong_para_a
    list_wrong_para_a_words = sorted(set())  # store List Paragraph style text that contain unacceptable spacing after paragraph values in the sorted list list_wrong_para_a_words
    CORRECT_SPACE_AFTER_LIST = None  # state the specified spacing after for List Paragraph style and store in the variable CORRECT_SPACE_AFTER_LIST; None = default = 8pt
    for paragraph in WordFile.paragraphs:
        if 'List Paragraph' == paragraph.style.name:
            for run in paragraph.runs:
                # add List Paragraph style text spacing after paragraph values in the set list_para_a
                list_para_a.add(paragraph.paragraph_format.space_after)
                # check if spacing after paragraph is unacceptable, if so, store in the set list_wrong_para_a
                if paragraph.paragraph_format.space_after is not CORRECT_SPACE_AFTER_LIST:
                    if paragraph.paragraph_format.space_after is not None:
                        list_wrong_para_a.add(paragraph.paragraph_format.space_after / 12700)
                    else:
                        list_wrong_para_a.add(paragraph.paragraph_format.space_after)
                    # append List Paragraph style text that contain unacceptable spacing after paragraph values in the sorted list list_wrong_para_a_words
                    list_wrong_para_a_words.append(run.text)

        # check if all elements in list_para_a are not CORRECT_SPACE_AFTER_LIST
    if list_para_a != {CORRECT_SPACE_AFTER_LIST} and len(list_para_a) != 0:
        # print this if all elements in list_para_a are not CORRECT_SPACE_AFTER_LIST and print list_wrong_para_a and list_wrong_para_a_words content
        st.error(f'''
        ❌ List Paragraph style text has incorrect spacing after paragraph: {', '.join(map(str, list_wrong_para_a))}  
        🡆 Incorrect spacing after paragraph for List Paragraph style text found here: {' >> '.join(map(str, list_wrong_para_a_words))}
        ''')
        # check if list_para_a is empty, if so this means that List Paragraph style was not found
    elif len(list_para_a) == 0:
        # print this if list_para_a is empty, since List Paragraph style was not found
        st.info("ℹ️List Paragraph style text spacing after paragraph not found as this style was not used.")

    # TODO line spacing code for Caption style text
    # Check line spacing for all text formatted in the Caption style (figure and table captions) ***************************
    cap_para_l = set()  # store all Caption style text line spacing values in the set cap_para_l
    cap_wrong_para_l = set()  # store unacceptable Caption style text line spacing values in the set cap_wrong_para_l
    cap_wrong_para_l_words = sorted(set())  # store Caption style text with unacceptable line spacing in the sorted list cap_wrong_para_l_words
    CORRECT_LINE_SPACE_CAP = None  # state the specified line spacing for Caption style and store in the variable CORRECT_LINE_SPACE_CAP; None = default = 1.0
    for paragraph in WordFile.paragraphs:
        if 'Caption' == paragraph.style.name:
            for run in paragraph.runs:
                # add Caption style text line spacing values in the set cap_para_l
                cap_para_l.add(paragraph.paragraph_format.line_spacing)
                # check if line spacing is unacceptable, if so, store in the set cap_wrong_para_l
                if paragraph.paragraph_format.line_spacing is not CORRECT_LINE_SPACE_CAP:
                    cap_wrong_para_l.add(paragraph.paragraph_format.line_spacing)
                    # append Caption style text that contain unacceptable line spacing in the sorted list cap_wrong_para_l_words
                    cap_wrong_para_l_words.append(run.text)

        # check if all elements in cap_para_l are not CORRECT_LINE_SPACE_CAP
    if cap_para_l != {CORRECT_LINE_SPACE_CAP} and len(cap_para_l) != 0:
        # print this if all elements in cap_para_l are not CORRECT_LINE_SPACE_CAP and print cap_wrong_para_l and cap_wrong_para_l_words content
        st.error(f'''
        ❌ Caption style text has incorrect line spacing: {', '.join(map(str, cap_wrong_para_l))}  
        🡆 Incorrect line spacing in Caption style text found here: {' >> '.join(map(str, cap_wrong_para_l_words))}
        ''')
        # check if cap_para_l is empty, if so this means that Caption style was not found
    elif len(cap_para_l) == 0:
        # print this if cap_para_l is empty, since Caption style was not found
        st.info("ℹ️Caption style text line spacing not found as this style was not used.")


# reference counter program function
def ref_count():
    # TODO start of Reference Counter code --------------------
    # add reference counter program banner
    st.subheader("**Reference Counter**")
    # Access the Word document
    TextDoc_ieee = docx.process(uploaded_file)

    # Count the number of IEEE citations in the document ********************
    # Regex pattern to find IEEE style citations
    pattern_ieee = r'(\[\d+-\d+\]|\[\d+(,\d+)*\])'
    # Try to find matches (returned as an iterator of matches) for IEEE
    results_ieee = re.finditer(pattern_ieee, TextDoc_ieee)

    # Build a list with IEEE citations obtained by looping through matches
    # Each match has the first and last indices of the match, relative to the original string
    references_ieee = [TextDoc_ieee[match.start(): match.end()] for match in results_ieee]
    # Remove duplicate citations
    unique_citations = list(set(references_ieee))
    # store number of unique citations as an integer in cite_num_int
    cite_num_int = len(unique_citations)
    # store number of unique citations as a string in cite_num_str
    cite_num_str = str(cite_num_int)

    # Check if the length of the list unique_citations is less than 5 and if it is not empty
    if len(unique_citations) < 5 and len(unique_citations) != 0:
        st.error(f'''
        ❌ Less than five IEEE style references were found.  
        🡆 Number of IEEE style references: {cite_num_str}
        ''')

    # Check if the list unique_citations is empty, meaning no IEEE references were found
    elif len(unique_citations) == 0:
        st.error("❌ No IEEE style references were found.")


# header and footer program function
def header_footer():
    # TODO start of Header and Footer code --------------------
    # add header and footer program banner
    st.subheader("**Header and Footer**")
    # access Word document file
    WordFile = Document(uploaded_file)

    # Extract headers and footers from document *******************
    headers = []  # store page headers in the list headers
    # loop through all sections in the document
    for section in WordFile.sections:
        for paragraph in section.header.paragraphs:
            # only append text from headers if it is not an empty string
            if '' != paragraph.text:
                headers.append(paragraph.text)

    # check if length of the list headers is not equal to zero, which implies headers exist in the document
    if len(headers) != 0:
        st.success(f"The following headers were found: {', '.join(map(str, headers))}")
    # check if length of the list headers is equal to zero, which implies headers do not exist in the document
    elif len(headers) == 0:
        st.error("❌ No headers were found.")

    footers = []  # store page footers in the list footers
    # loop through all sections in the document
    for section in WordFile.sections:
        for paragraph in section.footer.paragraphs:
            # only append text from footers if it is not an empty string
            if '' != paragraph.text:
                footers.append(paragraph.text)

    # check if length of the list footers is not equal to zero, which implies footers exist in the document
    if len(footers) != 0:
        st.success(f"The following footers were found: {', '.join(map(str, footers))}")
    # check if length of the list footers is equal to zero, which implies footers do not exist in the document
    elif len(footers) == 0:
        st.error("❌ No footers were found.")


# page margins program function
def page_margin():
    # TODO start of Page Margins code --------------------
    # add page margins program banner
    st.subheader("**Page Margins**")
    # access Word document file
    WordFile = Document(uploaded_file)

    # scan the document and check if page margins are valid ********************
    page_margins = []  # store page margins in the list page_margins
    # loop through all sections in the document
    for section in WordFile.sections:
        # append page margins in the list page_margins
        page_margins.append((section.top_margin, section.bottom_margin, section.left_margin, section.right_margin))

    # execute function to check if page margins correspond to those of Narrow margins
    # takes a value (margins) and returns true if it corresponds to those of Narrow margins
    def Narrow_present(margins):
        TOP_MARGIN = 457200  # top margin value for Narrow margins as interpreted by python-docx
        BOTTOM_MARGIN = 457200  # bottom margin value for Narrow margins as interpreted by python-docx
        LEFT_MARGIN = 457200  # left margin value for Narrow margins as interpreted by python-docx
        RIGHT_MARGIN = 457200  # right margin value for Narrow margins as interpreted by python-docx
        top, bottom, left, right = margins
        return top == TOP_MARGIN and bottom == BOTTOM_MARGIN and left == LEFT_MARGIN and right == RIGHT_MARGIN

    # obtain number of sections that do not contain Narrow margins and also convert them into strings
    no_Narrow = [str(i + 1) for i, section in enumerate(page_margins) if not Narrow_present(section)]

    # check if the length of the lists no_Narrow and page_margins are equal, if so this means that all pages do not have Narrow margins
    if len(no_Narrow) == len(page_margins):
        # print this if the length of the lists no_Narrow and page_margins are equal, meaning that all pages do not have Narrow margins
        st.error("❌ Whole document does not have Narrow margins.")

    # check if the list no_Narrow contains one element, if so this means that one section contains non-Narrow margins
    elif len(no_Narrow) == 1:
        # print this if the list no_Narrow contains one element, meaning that one section contains non-Narrow margins
        st.error(f"❌ Section {no_Narrow[0]} does not have Narrow margins.")

    # check if the list no_Narrow is not empty and if the length of the list page_margins is greater than the length of the list no_Narrow
    elif 0 < len(no_Narrow) < len(page_margins):
        st.error(f"❌ Sections {', '.join(no_Narrow)} do not have Narrow margins.")


# function to run selected programs(s)
def run_program():
    if cb2:
        font_name()
    if cb3:
        font_size()
    if cb4:
        font_colour()
    if cb5:
        char_form()
    if cb6:
        para_align()
    if cb7:
        spacing()
    if cb8:
        ref_count()
    if cb9:
        header_footer()
    if cb10:
        page_margin()


# configure sidebar text and widgets
st.sidebar.title("**Format Check v1.0**")
# select a document using file uploader
uploaded_file = st.sidebar.file_uploader("Choose a Word document", type='.docx', key=1)
cb1 = st.sidebar.checkbox('All', key=2)
# if 'All Formatting Items' is checked disable every other checkbox
if cb1:
    cb2 = st.sidebar.checkbox('Font Name', value=cb1, disabled=True, key=3)
    cb3 = st.sidebar.checkbox('Font Size', value=cb1, disabled=True, key=4)
    cb4 = st.sidebar.checkbox('Font Colour', value=cb1, disabled=True, key=5)
    cb5 = st.sidebar.checkbox('Character Formatting', value=cb1, disabled=True, key=6)
    cb6 = st.sidebar.checkbox('Paragraph Alignment', value=cb1, disabled=True, key=7)
    cb7 = st.sidebar.checkbox('Spacing', value=cb1, disabled=True, key=8)
    cb8 = st.sidebar.checkbox('Reference Counter', value=cb1, disabled=True, key=9)
    cb9 = st.sidebar.checkbox('Header and Footer', value=cb1, disabled=True, key=10)
    cb10 = st.sidebar.checkbox('Page Margins', value=cb1, disabled=True, key=11)
else:
    cb2 = st.sidebar.checkbox('Font Name', key=12)
    cb3 = st.sidebar.checkbox('Font Size', key=13)
    cb4 = st.sidebar.checkbox('Font Colour', key=14)
    cb5 = st.sidebar.checkbox('Character Formatting', key=15)
    cb6 = st.sidebar.checkbox('Paragraph Alignment', key=16)
    cb7 = st.sidebar.checkbox('Spacing', key=17)
    cb8 = st.sidebar.checkbox('Reference Counter', key=18)
    cb9 = st.sidebar.checkbox('Header and Footer', key=19)
    cb10 = st.sidebar.checkbox('Page Margins', key=20)

start_btn = st.sidebar.button('Start', on_click=run_program, key=21)
st.sidebar.text("")
st.sidebar.text("")

# style Start button
button_style = """
        <style>
        .stButton > button {
            color: black;
            font-size: 20px;
            background-color: #FFFFFF;
            border: 1px solid black;
            width: 305px;
            height: 50px;
        }
        </style>
        """
st.markdown(button_style, unsafe_allow_html=True)


# Hide hamburger menu and 'Made with Streamlit' footer
hide_streamlit_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            </style>
            """
st.markdown(hide_streamlit_style, unsafe_allow_html=True)
